"""Text extraction for ingest (FR-501).

Optional backends live in the ``memory`` extra: pypdfium2, python-docx,
trafilatura. md/txt need no extra packages.
"""

from __future__ import annotations

from pathlib import Path


def extract_text(path: Path) -> tuple[str, int]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        data = path.read_bytes()
        return data.decode("utf-8", errors="replace"), len(data)
    if suffix == ".html":
        return _extract_html(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(f"no extractor for {suffix}")


def _extract_html(path: Path) -> tuple[str, int]:
    try:
        import trafilatura
    except ImportError as exc:
        raise ImportError(
            "html ingest requires the memory extra: pip install -e '.[memory]'"
        ) from exc
    raw = path.read_bytes()
    text = trafilatura.extract(raw, include_comments=False, include_tables=True) or ""
    return text, len(raw)


def _extract_pdf(path: Path) -> tuple[str, int]:
    try:
        import pypdfium2 as pdfium  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ImportError(
            "pdf ingest requires the memory extra: pip install -e '.[memory]'"
        ) from exc
    data = path.read_bytes()
    document = pdfium.PdfDocument(data)
    pages: list[str] = []
    try:
        for index in range(len(document)):
            page = document[index]
            textpage = page.get_textpage()
            pages.append(textpage.get_text_bounded())
    finally:
        document.close()
    return "\n\n".join(pages), len(data)


def _extract_docx(path: Path) -> tuple[str, int]:
    try:
        from docx import Document as DocxDocument  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ImportError(
            "docx ingest requires the memory extra: pip install -e '.[memory]'"
        ) from exc
    data = path.read_bytes()
    document = DocxDocument(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs), len(data)
